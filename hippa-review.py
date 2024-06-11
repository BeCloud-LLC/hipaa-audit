# Written by James Phipps
# Edited by Serene Aryal on 05/29/2024

#modified scrip for api

import time
from datetime import datetime
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor

class HIPAAComplianceReview: 
    def __init__(self, customer_name, interviewer_name):
        self.customer_name = customer_name
        self.interviewer_name = interviewer_name
        self.sections_data = {}
        self.pillar = "HIPAA Compliance"

    def set_document_styles(self, doc):
        styles = {
            'Heading 1': {'name': 'Arial', 'size': Pt(16), 'bold': True},
            'Heading 2': {'name': 'Arial', 'size': Pt(14), 'bold': True},
            'Normal': {'name': 'Arial', 'size': Pt(11)}
        }
        for style_name, font_attrs in styles.items():
            style = doc.styles[style_name]
            style.font.name = font_attrs['name']
            style.font.size = font_attrs['size']
            style.font.bold = font_attrs.get('bold', False)

    def add_cover_page(self, doc):
        doc.add_paragraph(f'AWS {self.pillar} Pillar Review for {self.customer_name}\n', style='Title')
        doc.add_paragraph(f'Conducted by: {self.interviewer_name}\nDate: {datetime.now().strftime("%B %d, %Y")}', style='Normal')
        #doc.add_page_break()

    def add_header_logo(self, doc, logo_path):
        for section in doc.sections:
            header = section.header
            paragraph = header.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))


    def generate_dashboard(self, doc):
        needs_improvement = sum(1 for responses in self.sections_data.values() for r in responses if "not implemented" in r[1].lower() or r[3] == "alternative")
        total_responses = sum(len(responses) for responses in self.sections_data.values())
        high_risk_items = sum(1 for responses in self.sections_data.values() for r in responses if r[4] == "high")

        fig, axs = plt.subplots(1, 2, figsize=(12, 6))

        # Improvement Overview Pie Chart
        labels = 'Needs Improvement', 'No Improvement Needed'
        sizes = [needs_improvement, total_responses - needs_improvement]
        colors = ['gold', 'lightgreen']
        explode = (0.1, 0)
        axs[0].pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
        axs[0].axis('equal')
        axs[0].set_title("Improvement Overview")

        # High Risk Items Count Graph
        labels = ['High Risk Items', 'Other Items']
        sizes = [high_risk_items, total_responses - high_risk_items]
        colors = ['red', 'green']
        axs[1].bar(labels, sizes, color=colors)
        axs[1].set_title("High Risk Items Count")
        axs[1].set_ylabel('Number of Items')

        plt.tight_layout()
        img_data = BytesIO()
        plt.savefig(img_data, format='png')
        img_data.seek(0)

        doc.add_page_break()
        doc.add_heading("Dashboard", level=1)
        doc.add_picture(img_data, width=Inches(6))

    # def generate_immediate_action_plan(self, doc):
    #     doc.add_heading("Immediate Action Plan", level=1)
    #     doc.add_paragraph("The following action items are recommended to improve your HIPAA compliance:")

    #     sorted_data = sorted(
    #         ((section, q_text, rec, priority) for section, responses in self.sections_data.items() for q_text, _, rec, _, priority in responses if priority in ['high', 'medium', 'low']),
    #         key=lambda x: x[3], reverse=True
    #     )

    #     for section, q_text, rec, priority in sorted_data:
    #         doc.add_paragraph(f"{section} - {q_text}: {rec} (Priority: {priority.capitalize()})", style='List Bullet')

    def generate_immediate_action_plan(self, doc):
        doc.add_heading("Immediate Action Plan", level=1)
        doc.add_paragraph("The following action items are recommended to improve your HIPAA compliance:")

        sorted_data = sorted(
            ((priority, section, q_text, rec) for section, responses in self.sections_data.items() for q_text, _, rec, _, priority in responses if priority in ['high', 'medium', 'low']),
            key=lambda x: x[0], reverse=True
        )

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Priority'
        hdr_cells[1].text = 'Recommendation'
        hdr_cells[2].text = 'Section'

        for priority, section, q_text, rec in sorted_data:
            row_cells = table.add_row().cells
            row_cells[0].text = priority.capitalize()
            row_cells[1].text = f"{q_text}: {rec}"
            row_cells[2].text = section



    def collect_responses(self, sections, recommendations, hipaa_references):
        print(f"Starting review for {self.customer_name}.")
        print(f"Hello, {self.customer_name}. My name is {self.interviewer_name}, and I'm here to perform a HIPAA Compliance Review.")
        for section, questions in sections.items():
            section_responses = []
            print(f"\n--- Reviewing your {section} ---")
            for question_text, question_key in questions:
                # INPUT VALIDATION
                response = ""
                while response not in ["yes", "no", "n/a"]:
                    response = input(f"\n{question_text} (yes/no/n/a): ").strip().lower()
                    if response not in ["yes", "no", "n/a"]:
                        print("WRONG INPUT! The input should be yes, no or n/a")

                if response == "yes":
                    detailed_response = "Implemented"
                    recommendation = "No specific recommendation."
                    priority = ""
                elif response == "no":
                    detailed_response = "Not implemented"
                    recommendation = recommendations.get(question_key, "No specific recommendation.")

                    # INPUT VALIDATION
                    priority = ""
                    while priority not in ["high", "medium", "low"]:
                        priority = input(f"Here is our recommendation: {recommendation}\nHow would you rate the priority of this improvement? (high/medium/low): ").strip().lower()

                        if  priority not in ["high", "medium", "low"]:
                            print("WRONG INPUT! The input should be high, medium or low.")

                else:
                    continue
                hipaa_ref = hipaa_references.get(question_key, "No specific HIPAA reference.")
                section_responses.append((question_text, detailed_response, recommendation, hipaa_ref, priority))

            if section_responses:
                self.sections_data[section] = section_responses

    def create_document(self, start_time):
        doc = Document()
        self.set_document_styles(doc)
        self.add_cover_page(doc)
        self.add_header_logo(doc, "./logo.png")

        self.generate_dashboard(doc)
        self.generate_immediate_action_plan(doc)

        doc.add_heading("Detailed Responses", level=1)
        for section, responses in self.sections_data.items():
            doc.add_heading(section, level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Question'
            hdr_cells[1].text = 'Response'
            hdr_cells[2].text = 'Recommendation'
            hdr_cells[3].text = 'HIPAA Reference'

            for question_text, detailed_response, recommendation, hipaa_ref, _ in responses:
                row_cells = table.add_row().cells
                row_cells[0].text = question_text
                row_cells[1].text = detailed_response
                row_cells[2].text = recommendation
                row_cells[3].text = hipaa_ref

        duration = int(time.time() - start_time)
        filename = f"{self.customer_name}_{self.pillar}_Review_{datetime.now().strftime('%Y%m%d')}_Duration_{duration}s.docx"
        doc.save(filename)
        print(f"Document saved as '{filename}'.")

def main():
    customer_name = input("Enter the client's organization name: ")
    interviewer_name = input("Enter your name: ")

    pillar = "HIPAA Compliance"
    sections = {
        "AntiVirus Monitoring": [
            ("Do you have a managed AntiVirus solution in place for all devices?", "managed_antivirus"),
            ("Are regular AntiVirus updates and scans scheduled and monitored?", "regular_updates_scans"),
            ("Do you receive alerts for detected threats?", "threat_alerts"),
        ],
        "File Storage": [
            ("Do you use a secure and compliant file storage solution?", "secure_file_storage"),
            ("Is data encrypted at rest and in transit?", "data_encryption"),
            ("Do you have lifecycle policies for managing data retention?", "data_retention_policies"),
            ("Are regular backups conducted and stored securely?", "regular_backups"),
        ],
        "User Access Controls": [
            ("Do you have user access controls in place for data?", "user_access_controls"),
            ("Are there regular audits for user access and permissions?", "regular_audits"),
            ("Do you have multi-factor authentication (MFA) enabled?", "mfa_enabled"),
        ],
        "Firewalls and Web Filtering": [
            ("Are firewalls configured to protect against unauthorized access?", "firewall_configuration"),
            ("Is web filtering in place to prevent access to malicious websites?", "web_filtering"),
            ("Are Firewalls logs and alerts monitored for suspicious activity?", "logs_alerts_monitoring"),
        ],
        "VPNs": [
            ("Do you use VPNs to secure remote access to your network?", "use_vpns"),
            ("Are VPN connections encrypted?", "vpn_encryption"),
            ("Are VPN usage logs monitored for unusual activity?", "vpn_logs_monitoring"),
        ],
        "HIPAA Compliance": [
            ("Do you have documented policies and procedures for HIPAA compliance?", "hipaa_policies"),
            ("Are regular HIPAA compliance audits conducted?", "hipaa_audits"),
            ("Is staff trained regularly on HIPAA compliance?", "hipaa_training"),
        ],
        "Access Control": [
            ("Do you use unique user IDs for all employees accessing ePHI?", "unique_user_ids"),
            ("Do you have role-based access control (RBAC) in place?", "rbac"),
            ("Are access controls reviewed regularly?", "access_controls_review"),
        ],
        "Audit Controls": [
            ("Are audit logs maintained for all systems including ePHI?", "audit_logs"),
            ("Do you have automated mechanisms to review audit logs?", "automated_audit_review"),
            ("Are audit logs reviewed regularly?", "audit_logs_review"),
        ],
        "Transmission Security": [
            ("Is all data including ePHI encrypted during transmission?", "transmission_encryption"),
            ("Do you use secure methods for data transmission (e.g. TLS, VPN)?", "secure_transmission_methods"),
            ("Are secure channels - encrypted - used for email communication of ePHI?", "secure_email"),
        ],
        "Physical Safeguards": [
            ("Are physical access controls in place to protect systems storing ePHI?", "physical_access_controls"),
            ("Are backup media stored securely?", "secure_backup_storage"),
            ("Are facility access logs maintained?", "facility_access_logs"),
        ],
        "Administrative Safeguards": [
            ("Do you have a designated HIPAA compliance officer?", "hipaa_compliance_officer"),
            ("Are regular risk assessments conducted?", "regular_risk_assessments"),
            ("Are incident response plans in place?", "incident_response_plans"),
        ],
        "Breach Notification": [
            ("Do you have a documented breach notification policy?", "breach_notification_policy"),
            ("Are breaches reported to affected individuals within 60 days?", "breach_reporting"),
            ("Are breaches documented and reviewed to prevent recurrence?", "breach_documentation_review"),
        ],
        "Business Associate Agreements": [
            ("Do you have BAAs with all third parties accessing ePHI?", "baas"),
            ("Are BAAs reviewed and updated regularly?", "baas_review"),
            ("Are third parties monitored for compliance with BAAs?", "third_party_monitoring"),
        ],
    }

    actionable_recommendations = {
        "managed_antivirus": "Implement a managed AntiVirus solution like AWS Security Hub for all devices.",
        "regular_updates_scans": "Schedule regular updates and scans using BitDefender.",
        "threat_alerts": "Set up alerts for detected threats using Amazon GuardDuty.",
        "secure_file_storage": "Use a secure and compliant file storage solution like AWS which offers encryption at rest and in transit.",
        "data_encryption": "Ensure all sensitive data is encrypted at rest and in transit using AWS KMS or Drive Encryption.",
        "data_retention_policies": "Implement lifecycle policies using Amazon S3 for data retention management.",
        "regular_backups": "Conduct regular backups using AWS Backup and store them securely.",
        "user_access_controls": "Implement user access controls using AWS IAM or Active Directory for seamless identity management.",
        "regular_audits": "Perform regular audits of user access and permissions using AWS IAM Access Analyzer.",
        "mfa_enabled": "Enable multi-factor authentication (MFA) for added security.",
        "firewall_configuration": "Configure firewalls using AWS Network Firewall to protect against unauthorized access.",
        "web_filtering": "Implement web filtering using AWS WAF to prevent access to malicious websites.",
        "logs_alerts_monitoring": "Monitor logs and alerts for suspicious activity using Amazon CloudWatch.",
        "use_vpns": "Use VPNs like AWS Client VPN to secure remote access to your network.",
        "vpn_encryption": "Ensure VPN connections are encrypted using strong encryption protocols.",
        "vpn_logs_monitoring": "Monitor VPN usage logs for unusual activity using Amazon CloudWatch.",
        "hipaa_policies": "Document and regularly update policies and procedures for HIPAA compliance.",
        "hipaa_audits": "Conduct regular HIPAA compliance audits using AWS Audit Manager.",
        "hipaa_training": "Provide regular HIPAA compliance training to staff using AWS Training and Certification.",
        "unique_user_ids": "Implement unique user IDs for all employees accessing ePHI to ensure accountability.",
        "rbac": "Establish role-based access control (RBAC) to restrict access to ePHI based on job roles.",
        "access_controls_review": "Conduct regular reviews of access controls to ensure they are up-to-date and effective.",
        "audit_logs": "Maintain detailed audit logs for all systems accessing ePHI to ensure traceability.",
        "automated_audit_review": "Implement automated mechanisms to review audit logs for suspicious activities.",
        "audit_logs_review": "Review audit logs regularly to detect and respond to potential security incidents.",
        "transmission_encryption": "Encrypt all ePHI during transmission using strong encryption standards (e.g., AES-256).",
        "secure_transmission_methods": "Use secure methods like TLS and VPNs for data transmission to protect ePHI.",
        "secure_email": "Ensure secure channels (e.g., encrypted email services) are used for emailing ePHI.",
        "physical_access_controls": "Implement physical access controls (e.g., keycards, biometric scanners) to secure areas storing ePHI.",
        "secure_backup_storage": "Store backup media in secure access-controlled environments.",
        "facility_access_logs": "Maintain facility access logs to track and monitor physical access to sensitive areas.",
        "hipaa_compliance_officer": "Designate a HIPAA compliance officer responsible for overseeing HIPAA compliance efforts.",
        "regular_risk_assessments": "Conduct regular risk assessments to identify and mitigate potential HIPAA compliance risks.",
        "incident_response_plans": "Develop and maintain incident response plans to handle potential HIPAA breaches.",
        "breach_notification_policy": "Document a breach notification policy that outlines steps to be taken in the event of a HIPAA breach.",
        "breach_reporting": "Ensure breaches are reported to affected individuals within 60 days as required by HIPAA.",
        "breach_documentation_review": "Document and review breaches to identify root causes and prevent recurrence.",
        "baas": "Establish Business Associate Agreements (BAAs) with all third parties accessing ePHI.",
        "baas_review": "Review and update BAAs regularly to ensure they meet current HIPAA requirements.",
        "third_party_monitoring": "Monitor third parties for compliance with BAAs and take corrective actions if necessary.",
    }

    hipaa_references = {
        "managed_antivirus": "HIPAA Security Rule: 45 CFR §164.308(a)(5)(ii)(B)",
        "regular_updates_scans": "HIPAA Security Rule: 45 CFR §164.308(a)(5)(ii)(B)",
        "threat_alerts": "HIPAA Security Rule: 45 CFR §164.308(a)(5)(ii)(B)",
        "secure_file_storage": "HIPAA Security Rule: 45 CFR §164.310(d)(1)",
        "data_encryption": "HIPAA Security Rule: 45 CFR §164.312(e)(2)(B)",
        "data_retention_policies": "HIPAA Security Rule: 45 CFR §164.310(d)(2)(i)",
        "regular_backups": "HIPAA Security Rule: 45 CFR §164.308(a)(7)(ii)(B)",
        "user_access_controls": "HIPAA Security Rule: 45 CFR §164.312(a)(1)",
        "regular_audits": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(D)",
        "mfa_enabled": "HIPAA Security Rule: 45 CFR §164.312(d)",
        "firewall_configuration": "HIPAA Security Rule: 45 CFR §164.312(c)(1)",
        "web_filtering": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(B)",
        "logs_alerts_monitoring": "HIPAA Security Rule: 45 CFR §164.312(b)",
        "use_vpns": "HIPAA Security Rule: 45 CFR §164.312(e)(1)",
        "vpn_encryption": "HIPAA Security Rule: 45 CFR §164.312(e)(2)(B)",
        "vpn_logs_monitoring": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(D)",
        "hipaa_policies": "HIPAA Security Rule: 45 CFR §164.316(a)",
        "hipaa_audits": "HIPAA Security Rule: 45 CFR §164.308(a)(8)",
        "hipaa_training": "HIPAA Security Rule: 45 CFR §164.308(a)(5)",
        "unique_user_ids": "HIPAA Security Rule: 45 CFR §164.312(a)(2)(i)",
        "rbac": "HIPAA Security Rule: 45 CFR §164.312(a)(1)",
        "access_controls_review": "HIPAA Security Rule: 45 CFR §164.312(a)(1)",
        "audit_logs": "HIPAA Security Rule: 45 CFR §164.312(b)",
        "automated_audit_review": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(D)",
        "audit_logs_review": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(D)",
        "transmission_encryption": "HIPAA Security Rule: 45 CFR §164.312(e)(1)",
        "secure_transmission_methods": "HIPAA Security Rule: 45 CFR §164.312(e)(2)(B)",
        "secure_email": "HIPAA Security Rule: 45 CFR §164.312(e)(1)",
        "physical_access_controls": "HIPAA Security Rule: 45 CFR §164.310(a)(1)",
        "secure_backup_storage": "HIPAA Security Rule: 45 CFR §164.310(d)(1)",
        "facility_access_logs": "HIPAA Security Rule: 45 CFR §164.310(a)(2)(ii)",
        "hipaa_compliance_officer": "HIPAA Security Rule: 45 CFR §164.308(a)(2)",
        "regular_risk_assessments": "HIPAA Security Rule: 45 CFR §164.308(a)(1)(ii)(A)",
        "incident_response_plans": "HIPAA Security Rule: 45 CFR §164.308(a)(6)(ii)",
        "breach_notification_policy": "HIPAA Breach Notification Rule: 45 CFR §§164.400-414",
        "breach_reporting": "HIPAA Breach Notification Rule: 45 CFR §§164.404",
        "breach_documentation_review": "HIPAA Breach Notification Rule: 45 CFR §§164.408",
        "baas": "HIPAA Privacy Rule: 45 CFR §164.502(e)",
        "baas_review": "HIPAA Privacy Rule: 45 CFR §164.502(e)(2)",
        "third_party_monitoring": "HIPAA Privacy Rule: 45 CFR §164.504(e)(1)(ii)",
    }

    review = HIPAAComplianceReview(customer_name, interviewer_name)
    review.collect_responses(sections, actionable_recommendations, hipaa_references)

    start_time = time.time()
    review.create_document(start_time)

if __name__ == "__main__":
    main()
